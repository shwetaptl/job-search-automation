"""
Job Search Automation

Usage:
  python job_search.py --init path/to/profile.docx   build profile_data.md from your resume doc
  python job_search.py --update "fact to add"         merge a new fact into profile_data.md
  python job_search.py --show                         print the current profile_data.md
  python job_search.py                                run the job search using profile_data.md

Outputs: jobs_YYYYMMDD_HHMMSS.csv sorted by match score (highest first)
"""

import sys
import csv
import re
import time
import os
import argparse
from datetime import datetime

import requests
from bs4 import BeautifulSoup
from docx import Document
try:
    from playwright.sync_api import sync_playwright
    PLAYWRIGHT_AVAILABLE = True
except ImportError:
    PLAYWRIGHT_AVAILABLE = False

# ── Config ────────────────────────────────────────────────────────────────────

GITHUB_SOURCES = [
    {
        "name": "speedyapply",
        "raw_url": "https://raw.githubusercontent.com/speedyapply/2026-SWE-College-Jobs/main/NEW_GRAD_USA.md",
        "format": "markdown-html",
    },
    {
        "name": "SimplifyJobs",
        "raw_url": "https://raw.githubusercontent.com/SimplifyJobs/New-Grad-Positions/dev/README.md",
        "format": "html-table",
    },
    {
        "name": "jobright-ai",
        "raw_url": "https://raw.githubusercontent.com/jobright-ai/2026-Software-Engineer-New-Grad/master/README.md",
        "format": "markdown-html",
    },
    {
        "name": "cvrve",
        "raw_url": "https://raw.githubusercontent.com/cvrve/New-Grad/main/README.md",
        "format": "markdown-html",
    },
]

# Companies that publish open job boards via the Greenhouse API.
# Add or remove slugs freely — the API is public and requires no key.
GREENHOUSE_COMPANIES = [
    "inovalon",  # Inovalon (59 jobs)
    "athena",  # athenahealth (2 jobs)
    "brex",  # Brex (252 jobs)
    "robinhood",  # Robinhood (132 jobs)
    "coinbase",  # Coinbase (129 jobs)
    "sofi",  # SoFi (93 jobs)
    "affirm",  # Affirm (172 jobs)
    "marqeta",  # Marqeta (35 jobs)
    "chime",  # Chime (64 jobs)
    "block",  # Square / Block (206 jobs)
    "solera",  # Solera (0 jobs)
    "nice",  # NICE Systems (338 jobs)
    "qualtrics",  # Qualtrics (59 jobs)
    "hubspot",  # HubSpot (0 jobs)
    "okta",  # Okta (374 jobs)
    "mongodb",  # MongoDB (400 jobs)
    "elastic",  # Elastic (205 jobs)
    "datadog",  # Datadog (412 jobs)
    "twilio",  # Twilio (152 jobs)
    "cloudflare",  # Cloudflare (228 jobs)
    "zscaler",  # Zscaler (317 jobs)
    "databricks",  # Databricks (786 jobs)
    "newrelic",  # New Relic (55 jobs)
    "pagerduty",  # PagerDuty (24 jobs)
    "sumologic",  # Sumo Logic (24 jobs)
    "grafanalabs",  # Grafana Labs (99 jobs)
    "sisense",  # Sisense (11 jobs)
    "amplitude",  # Amplitude (49 jobs)
    "lattice",  # Lattice (11 jobs)
    "discord",  # Discord (62 jobs)
    "reddit",  # Reddit (183 jobs)
    "asana",  # Asana (151 jobs)
    "flexport",  # Flexport (129 jobs)
    "pingidentity",  # Ping Identity (37 jobs)
    "onespan",  # OneSpan (28 jobs)
    "thoughtworks",  # Thoughtworks (79 jobs)
    "tcs",  # TCS (77 jobs)
    "stripe",  # Stripe (488 jobs)
    "airbnb",  # Airbnb (225 jobs)
    "figma",  # Figma (172 jobs)
    "lyft",  # Lyft (151 jobs)
    "pinterest",  # Pinterest (187 jobs)
    "twitch",  # Twitch (64 jobs)
    "dropbox",  # Dropbox (55 jobs)
    "gusto",  # Gusto (78 jobs)
    "anthropic",  # Anthropic (389 jobs)
    "harnessinc",  # was: harness             # Harness — CI/CD platform
    "sentinellabs",        # SentinelOne — cybersecurity
    "fractalsoftware",  # was: fractal             # Fractal Analytics — AI/data
]

# Companies using Lever ATS — verified via API, free public JSON, full descriptions in response.
LEVER_COMPANIES = [
    "ro",             # healthcare tech — backend SWE roles
    "metabase",       # data/analytics — Python/Java backend
    "outreach",       # sales tech SaaS — backend eng
    "pivotal",        # software consulting — Java/Python
    "mistral",        # AI/LLM — Paris-based, H1B unclear but strong eng brand
    "cloudinary",     # media CDN — backend/SDK eng
    "zimperium",      # mobile cybersecurity
    "anomali",        # threat intelligence cybersecurity
    "coalfire",       # cybersecurity consulting
    "contentsquare",  # digital analytics
    "nium",           # fintech payments — H1B sponsor
    "matillion",      # data integration/ETL
    "highspot",       # sales enablement SaaS
    "arcadia",        # healthcare data platform
    "sonatype",       # dev tools / supply chain security
    "logrocket",      # frontend observability
    "neon",               # Postgres cloud database
    "plaid",              # fintech payments — moved from Greenhouse
    "pointclickcare",     # senior care SaaS — H1B sponsor
]

# Companies using Workday ATS — search API returns titles/URLs; descriptions fetched via Playwright.
# Format: (tenant_slug, career_board_name)
WORKDAY_COMPANIES = [
    ("servicenow", "ServiceNow"),
    ("snowflake", "Snowflake"),
    ("crowdstrike", "crowdstrike"),
    ("medtronic", "medtronic"),
    ("gehealthcare", "GE_HealthCare"),
    ("philips", "philips"),
    ("tylertech", "Tyler_Technologies"),
    ("trimble", "Trimble"),
    ("fiserv", "fiserv"),
    ("fisglobal", "FIS"),
    ("jackhenry", "jackhenry"),
    ("broadridge", "Broadridge"),
    ("morningstar", "morningstar"),
    ("autodesk", "Autodesk"),
    ("paloaltonetworks", "paloaltonetworks"),
    ("fortinet", "fortinet"),
    ("sailpoint", "sailpoint"),
    ("rapid7", "rapid7"),
    ("dynatrace", "dynatrace"),
    ("splunk", "splunk"),
    ("procore", "procore"),
    ("atlassian", "atlassian"),
    ("instructure", "instructure"),
    ("opentext", "opentext"),
    ("pegasystems", "pega"),
    ("informatica", "informatica"),
    ("epic", "epic"),
    ("ssc", "ssc"),
    ("bentleysystems", "bentley"),
    ("ptc", "ptc"),
    ("verint", "verint"),
    ("infor", "infor"),
    ("epicor", "epicor"),
    ("realpage", "realpage"),
    ("yardi", "yardi"),
    ("mrisoftware", "mrisoftware"),
    ("tradeweb", "tradeweb"),
    ("envestnet", "envestnet"),
    ("wipro", "wipro"),
    ("infosys", "infosys"),
    ("hcltech", "hcltech"),
    ("capgemini", "capgemini"),
    ("ltimindtree", "ltimindtree"),
    ("cognizant", "cognizant"),
    ("quantiphi", "Careers_at_Quantiphi"),  # AI/data — moved from Greenhouse
    ("guidewire", "external"),              # insurance platform — wd5
]

# Companies where the ATS slug could not be auto-confirmed.
# To fix: google "{company name} careers site:myworkdayjobs.com" (Workday)
#         or visit boards.greenhouse.io/{slug} (Greenhouse)
#         or visit jobs.lever.co/{slug} (Lever)
# BEGIN_UNCONFIRMED
UNCONFIRMED_SLUGS = {
    # "Company Name": ("ats_type", "slug_tried", "manual_fix_url"),
}
# END_UNCONFIRMED

# Titles containing these words are skipped — they signal senior/management roles
# not suitable for a new-grad candidate.
SENIOR_TITLE_KEYWORDS = [
    "senior", "staff", "principal", "manager", "director", "lead",
    "head of", "vp ", "vice president", "distinguished", "fellow",
    "architect", "cto", "ceo", "ciso",
]

GEMINI_MODEL = "gemini-2.0-flash"
GROQ_MODEL   = "llama-3.3-70b-versatile"
MIN_SCORE = 50          # jobs below this score are skipped in output
MAX_JOBS = 200          # safety cap — process at most this many jobs
REQUEST_DELAY = 0.5     # seconds between HTTP requests (be polite)
REQUEST_TIMEOUT = 10    # seconds per HTTP request

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/120.0.0.0 Safari/537.36"
    )
}

PROFILE_DOC_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "profile_data.md")

# ── Helpers ───────────────────────────────────────────────────────────────────

def read_docx(path: str) -> str:
    """Extract all text from a .docx file."""
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # also grab tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def summarize_profile(provider: str, client, raw_text: str) -> str:
    """
    One-time call: ask Gemini to distill the raw .docx text into a clean,
    scoring-optimized summary. Strips meta-instructions, interview coaching
    notes, and conditional resume rules — keeps only factual career info.
    """
    prompt = f"""You are preparing a candidate profile for automated job matching against software engineering job descriptions.

Below is a detailed career profile document. Extract ONLY the factual career information relevant to job matching.

EXCLUDE:
- Meta-instructions or notes about how to use the document
- Interview coaching advice
- Conditional rules about resume writing (e.g. "only list X if JD requires it")
- "Do not claim" / "not skilled in" disclaimers (just omit those skills entirely)
- Internal notes about what to say vs not say

INCLUDE with all details:
- Authorization status (visa type, OPT start date, STEM extension eligibility, relocation openness)
- Education (degrees, schools, GPAs, graduation dates)
- Target roles (what the candidate is open to)
- Work experience (company, title, dates, technologies used, key achievements and metrics)
- Personal projects (tech stack, what it does, quantified outcomes)
- Technical skills (verified skills only — languages, frameworks, cloud services, databases, tools)
- Certifications

Output a clean structured summary. Keep all real metrics and tech stack details — they are critical for accurate job matching.

PROFILE DOCUMENT:
{raw_text}"""

    try:
        return call_llm(provider, client, prompt, max_tokens=2048)
    except Exception as e:
        print(f"  [warn] Profile summarization failed: {e}. Using raw text.")
        return raw_text


def load_profile_doc() -> str:
    """Load the persistent profile doc. Errors out with guidance if it doesn't exist yet."""
    if not os.path.exists(PROFILE_DOC_PATH):
        print(f"Error: no profile doc found at {PROFILE_DOC_PATH}")
        print("Run this first: python job_search.py --init path/to/your_profile.docx")
        sys.exit(1)
    with open(PROFILE_DOC_PATH, "r", encoding="utf-8") as f:
        return f.read()


def save_profile_doc(text: str) -> None:
    """Write the persistent profile doc to disk."""
    with open(PROFILE_DOC_PATH, "w", encoding="utf-8") as f:
        f.write(text)


def update_profile_doc(provider: str, client, current_doc: str, new_fact: str) -> str:
    """
    Merge a new fact (e.g. "I learned Kubernetes", "I completed an AWS cert")
    into the existing structured profile doc. Keeps everything else unchanged
    unless the new fact directly contradicts something already there.
    """
    prompt = f"""You maintain a structured career profile document used for automated job matching.

Below is the CURRENT profile document, followed by a NEW FACT the candidate wants added.

Update the document to incorporate the new fact:
- Add it to the most relevant existing section (e.g. a new skill goes under Technical Skills, a new project goes under Projects)
- If the new fact contradicts something already in the document, update that part rather than duplicating it
- Keep everything else in the document unchanged
- Keep the same structured format (same section headers) as the current document
- Do not add commentary, meta-notes, or explanations — output only the updated profile document

CURRENT PROFILE DOCUMENT:
{current_doc}

NEW FACT TO ADD:
{new_fact}

Output the complete updated profile document:"""

    return call_llm(provider, client, prompt, max_tokens=2048)


def fetch_text(url: str) -> str:
    """Fetch URL and return raw text. Returns empty string on failure."""
    try:
        r = requests.get(url, headers=HEADERS, timeout=REQUEST_TIMEOUT)
        r.raise_for_status()
        return r.text
    except Exception as e:
        print(f"  [warn] Could not fetch {url}: {e}")
        return ""


def looks_js_rendered(html: str) -> bool:
    """Return True if the page HTML looks like a JS-rendered shell with no real content."""
    markers = [
        "__NEXT_DATA__", "ng-version", "data-reactroot",
        '<div id="root">', '<div id="app">', "window.__INITIAL_STATE__",
        "__nuxt__", "data-v-app",
    ]
    return any(m in html for m in markers)


def fetch_with_browser(url: str) -> str:
    """Fetch a URL using a headless Chromium browser and return the rendered HTML."""
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(url, timeout=20000, wait_until="domcontentloaded")
            page.wait_for_timeout(2500)  # let JS render
            html = page.content()
            browser.close()
            return html
    except Exception as e:
        print(f"  [warn] Browser fetch failed for {url}: {e}")
        return ""


def _parse_html_to_text(html: str) -> str:
    """Shared BeautifulSoup extraction used by both fetch paths."""
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "header", "footer", "aside"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    return "\n".join(lines)


def extract_job_description(url: str) -> str:
    """
    Fetch a job posting and extract visible text.
    Tries plain HTTP first; falls back to headless Chromium if the page
    looks JS-rendered (empty shell returned by requests).
    """
    html = fetch_text(url)
    if not html:
        return ""

    text = _parse_html_to_text(html)

    if len(text) < 200 and looks_js_rendered(html):
        if PLAYWRIGHT_AVAILABLE:
            print("  [browser] JS-rendered page detected, using headless browser...")
            html = fetch_with_browser(url)
            if html:
                text = _parse_html_to_text(html)
        else:
            print("  [warn] JS-rendered page detected but Playwright not installed.")
            print("         Run: pip install playwright && playwright install chromium")

    return text[:8000]


def _extract_urls(text: str) -> list[str]:
    """Extract all https URLs from text (markdown or HTML link format)."""
    urls = []
    urls += re.findall(r'href=["\']?(https?://[^"\'>\s]+)', text)
    urls += [u for _, u in re.findall(r"\[([^\]]*)\]\((https?://[^)]+)\)", text)]
    return urls


def _clean_text(text: str) -> str:
    """Strip HTML and markdown link syntax, keep visible text only."""
    text = re.sub(r"<a[^>]*>(.*?)</a>", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"</?(?:strong|em|b|i|div|td|tr|th|tbody|thead)[^>]*>", "", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\[([^\]]*)\]\([^)]*\)", r"\1", text)
    text = re.sub(r":[a-z_]+:", "", text)  # :emoji: shortcodes
    text = re.sub(r"\*+([^*]+)\*+", r"\1", text)  # **bold** and *italic*
    return text.strip()


def _pick_apply_url(urls: list[str]) -> str:
    """
    From a list of URLs in a job row, pick the best apply link.
    Prefer non-github, non-company-homepage URLs (i.e. ATS / job board links).
    Simplify links are good apply proxies if nothing better is found.
    """
    ats_keywords = (
        "myworkday", "greenhouse", "lever", "ashby", "smartrecruiters",
        "icims", "jobvite", "workday", "taleo", "ultipro", "bamboohr",
        "amazon.jobs", "careers.", "jobs.", "job-boards", "apply",
        "simplify.jobs/p/",  # simplify direct apply pages
    )
    # tier 1: real ATS links
    for url in urls:
        if any(kw in url.lower() for kw in ats_keywords):
            return url
    # tier 2: any non-github link
    for url in urls:
        if "github.com" not in url and "simplify.jobs/c/" not in url:
            return url
    # tier 3: simplify company page (still useful)
    for url in urls:
        if "simplify.jobs" in url:
            return url
    return urls[-1] if urls else ""


def is_posted_today(posted: str) -> bool:
    """
    Decide whether a job's 'posted' string means it went up within the last 24 hours.
    Handles two formats seen across the GitHub source repos:
      - relative age: "0d", "5h", "30m", "new" -> today
      - absolute date: "Jun 30" -> compare month/day to today's date
    Unparseable/missing values are treated as NOT today (conservative).
    """
    if not posted:
        return False
    p = posted.strip().lower()

    if p in ("new", "today"):
        return True

    m = re.match(r"^(\d+)\s*([hmd])\b", p)
    if m:
        value, unit = int(m.group(1)), m.group(2)
        if unit == "d":
            return value == 0
        return True  # hours or minutes ago is always within 24h

    # ISO timestamp from Greenhouse: "2026-06-30T17:05:44-04:00"
    try:
        parsed = datetime.fromisoformat(posted.strip())
        today = datetime.now()
        return parsed.date() == today.date()
    except ValueError:
        pass

    # absolute date like "Jun 30"
    try:
        parsed = datetime.strptime(posted.strip(), "%b %d")
        today = datetime.now()
        return parsed.month == today.month and parsed.day == today.day
    except ValueError:
        return False


def parse_jobs(content: str, fmt: str) -> list[dict]:
    """
    Dispatch to the right parser based on fmt:
      'markdown-html' — pipe-delimited markdown table with optional HTML in cells
      'html-table'    — raw HTML <table><tr><td> embedded in the file
    """
    if fmt == "html-table":
        return _parse_html_table(content)
    return _parse_markdown_table(content)


def _parse_markdown_table(markdown: str) -> list[dict]:
    jobs = []
    for line in markdown.splitlines():
        if not line.startswith("|"):
            continue
        if re.match(r"^\|[\s\-:|]+\|$", line):
            continue
        cells = [c.strip() for c in line.split("|")]
        cells = [c for c in cells if c]
        if len(cells) < 3:
            continue

        all_urls = []
        for cell in cells:
            all_urls.extend(_extract_urls(cell))

        if not all_urls:
            continue

        apply_url = _pick_apply_url(all_urls)
        if not apply_url:
            continue

        company = _clean_text(cells[0])
        role = _clean_text(cells[1]) if len(cells) > 1 else ""
        location = _clean_text(cells[2]) if len(cells) > 2 else ""
        posted = _clean_text(cells[-1]) if len(cells) > 3 else ""

        if company.lower() in ("company", "name", ""):
            continue

        jobs.append({"company": company, "role": role, "location": location, "url": apply_url, "posted": posted})

    return jobs


def _parse_html_table(content: str) -> list[dict]:
    """Parse jobs from HTML <tr><td> tables embedded in the file."""
    soup = BeautifulSoup(content, "html.parser")
    jobs = []
    for row in soup.find_all("tr"):
        cells = row.find_all("td")
        if len(cells) < 3:
            continue

        # column 0 = company, 1 = role, 2 = location, 3 = apply link(s), last = age
        company = _clean_text(cells[0].get_text())
        role = _clean_text(cells[1].get_text())
        location = _clean_text(cells[2].get_text())
        posted = _clean_text(cells[-1].get_text()) if len(cells) > 3 else ""

        if company.lower() in ("company", "name", ""):
            continue

        # gather all URLs from the row
        all_urls = _extract_urls(str(row))
        if not all_urls:
            continue

        apply_url = _pick_apply_url(all_urls)
        if not apply_url:
            continue

        jobs.append({"company": company, "role": role, "location": location, "url": apply_url, "posted": posted})

    return jobs


# ─── Slug validation helpers ──────────────────────────────────────────────────

def _check_greenhouse(slug: str) -> int | None:
    """Return job count if slug returns HTTP 200, else None."""
    url = f"https://boards-api.greenhouse.io/v1/boards/{slug}/jobs"
    try:
        r = requests.get(url, headers=HEADERS, timeout=REQUEST_TIMEOUT)
        if r.status_code == 200:
            return len(r.json().get("jobs", []))
    except Exception:
        pass
    return None


def _check_lever(slug: str) -> int | None:
    """Return posting count if slug returns HTTP 200 with a JSON list, else None."""
    url = f"https://api.lever.co/v0/postings/{slug}?mode=json"
    try:
        r = requests.get(url, headers=HEADERS, timeout=REQUEST_TIMEOUT)
        if r.status_code == 200:
            data = r.json()
            if isinstance(data, list):
                return len(data)
    except Exception:
        pass
    return None


def _greenhouse_variations(slug: str) -> list[str]:
    """Return candidate slug variations to try when base slug fails."""
    shorthand = {
        "modernizingmedicine": "modmed", "pointclickcare": "pcc",
        "guidewiresoftware": "guidewire", "duckcreektechnologies": "duckcreek",
        "appliedsystems": "ezlynx", "westmonroe": "westmonroepartners",
    }
    candidates = []
    stripped = re.sub(r"[^a-z0-9]", "", slug)
    if stripped != slug:
        candidates.append(stripped)
    candidates.append(slug + "inc")
    for suffix in ["technologies", "systems", "software", "labs", "health"]:
        v = slug + suffix
        if v != slug:
            candidates.append(v)
    if slug in shorthand:
        candidates.append(shorthand[slug])
    seen: set[str] = {slug}
    return [c for c in candidates if c not in seen and not seen.add(c)]  # type: ignore[func-returns-value]


def _lever_variations(slug: str) -> list[str]:
    """Return candidate Lever slug variations to try when base slug fails."""
    shorthand = {
        "guidewiresoftware": "guidewire", "duckcreektechnologies": "duckcreek",
        "publicissapient": "sapient", "netsmart": "netsmartechnologies",
    }
    candidates = []
    no_hyphens = slug.replace("-", "")
    if no_hyphens != slug:
        candidates.append(no_hyphens)
    for suffix in ["inc", "hq", "careers", "software"]:
        candidates.append(slug + suffix)
    for suffix in ["software", "inc", "hq"]:
        if slug.endswith(suffix) and len(slug) > len(suffix):
            candidates.append(slug[: -len(suffix)])
    if slug in shorthand:
        candidates.append(shorthand[slug])
    seen: set[str] = {slug}
    return [c for c in candidates if c not in seen and not seen.add(c)]  # type: ignore[func-returns-value]


def _workday_tenant_variations(tenant: str) -> list[str]:
    """Return candidate Workday tenant slug variations to try when base tenant fails."""
    shorthand = {
        "ltimindtree": "mindtree", "mrisoftware": "mri",
        "fisglobal": "fis", "bentleysystems": "bentley",
    }
    candidates = []
    for suffix in ["inc", "corp", "llc"]:
        candidates.append(tenant + suffix)
    for suffix in ["inc", "corp", "systems", "technologies"]:
        if tenant.endswith(suffix) and len(tenant) > len(suffix):
            candidates.append(tenant[: -len(suffix)])
    candidates.append(tenant + "careers")
    if tenant in shorthand:
        candidates.append(shorthand[tenant])
    seen: set[str] = {tenant}
    return [c for c in candidates if c not in seen and not seen.add(c)]  # type: ignore[func-returns-value]


def _patch_job_search_py(results: dict) -> None:
    """Rewrite job_search.py in-place: fix confirmed slugs, comment out failures, update UNCONFIRMED_SLUGS."""
    script_path = os.path.abspath(__file__)
    with open(script_path, "r") as f:
        content = f.read()

    failed_entries = []

    for key, info in results.items():
        ats = info["ats"]
        status = info["status"]
        fixed = info.get("fixed")

        if ats == "greenhouse":
            if status == "fixed":
                content = re.sub(
                    rf'(\s+)"{re.escape(key)}"(,)',
                    rf'\1"{fixed}"\2  # was: {key}',
                    content, count=1,
                )
            elif status == "failed":
                content = re.sub(
                    rf'^(\s+)"{re.escape(key)}"(,.*)',
                    rf'\1# "{key}"\2  # UNCONFIRMED',
                    content, count=1, flags=re.MULTILINE,
                )
                failed_entries.append((key.title(), "greenhouse", key,
                                       f"https://boards.greenhouse.io/{key}"))

        elif ats == "lever":
            if status == "fixed":
                content = re.sub(
                    rf'(\s+)"{re.escape(key)}"(,)',
                    rf'\1"{fixed}"\2  # was: {key}',
                    content, count=1,
                )
            elif status == "failed":
                content = re.sub(
                    rf'^(\s+)"{re.escape(key)}"(,.*)',
                    rf'\1# "{key}"\2  # UNCONFIRMED',
                    content, count=1, flags=re.MULTILINE,
                )
                failed_entries.append((key.title(), "lever", key,
                                       f"https://jobs.lever.co/{key}"))

    # Rebuild UNCONFIRMED_SLUGS block using sentinel comments as anchors
    if failed_entries:
        new_entries = "\n".join(
            f'    "{name}": ("{ats}", "{slug}", "{url}"),'
            for name, ats, slug, url in failed_entries
        )
        new_block = (
            "# BEGIN_UNCONFIRMED\n"
            "UNCONFIRMED_SLUGS = {\n"
            '    # "Company Name": ("ats_type", "slug_tried", "manual_fix_url"),\n'
            f"{new_entries}\n"
            "}\n"
            "# END_UNCONFIRMED"
        )
        content = re.sub(
            r"# BEGIN_UNCONFIRMED\nUNCONFIRMED_SLUGS\s*=\s*\{.*?\}\n# END_UNCONFIRMED",
            new_block,
            content,
            flags=re.DOTALL,
        )

    with open(script_path, "w") as f:
        f.write(content)


def run_slug_validation() -> None:
    """Test every ATS slug, auto-try variations, patch job_search.py, print summary."""
    results = {}

    # ── Greenhouse ─────────────────────────────────────────────────────────────
    print("\nValidating Greenhouse slugs...")
    for slug in GREENHOUSE_COMPANIES:
        count = _check_greenhouse(slug)
        if count is not None:
            results[slug] = {"ats": "greenhouse", "status": "ok", "count": count}
            print(f"  ✅ greenhouse/{slug} → {count} jobs")
        else:
            fixed_slug = None
            fixed_count = 0
            for variant in _greenhouse_variations(slug):
                c = _check_greenhouse(variant)
                if c is not None:
                    fixed_slug = variant
                    fixed_count = c
                    break
            if fixed_slug:
                results[slug] = {"ats": "greenhouse", "status": "fixed",
                                 "count": fixed_count, "fixed": fixed_slug}
                print(f"  ✅ greenhouse/{slug} → fixed as '{fixed_slug}' ({fixed_count} jobs)")
            else:
                results[slug] = {"ats": "greenhouse", "status": "failed"}
                print(f"  ❌ greenhouse/{slug} → no working slug found")
        time.sleep(0.3)

    # ── Lever ──────────────────────────────────────────────────────────────────
    print("\nValidating Lever slugs...")
    for slug in LEVER_COMPANIES:
        count = _check_lever(slug)
        if count is not None:
            results[slug] = {"ats": "lever", "status": "ok", "count": count}
            print(f"  ✅ lever/{slug} → {count} postings")
        else:
            fixed_slug = None
            fixed_count = 0
            for variant in _lever_variations(slug):
                c = _check_lever(variant)
                if c is not None:
                    fixed_slug = variant
                    fixed_count = c
                    break
            if fixed_slug:
                results[slug] = {"ats": "lever", "status": "fixed",
                                 "count": fixed_count, "fixed": fixed_slug}
                print(f"  ✅ lever/{slug} → fixed as '{fixed_slug}' ({fixed_count} postings)")
            else:
                results[slug] = {"ats": "lever", "status": "failed"}
                print(f"  ❌ lever/{slug} → no working slug found")
        time.sleep(0.3)

    # ── Workday ────────────────────────────────────────────────────────────────
    # Workday uses wildcard DNS + Cloudflare CDN — HTTP-based validation returns the
    # same 422 error for all tenants (real or fake) without a real browser session.
    # Validation requires Playwright and is deferred to the actual fetch run.
    print(f"\nWorkday: HTTP validation not possible (Cloudflare-protected, wildcard DNS).")
    print(f"  Skipping {len(WORKDAY_COMPANIES)} Workday tenants — validated at runtime by fetch_workday_jobs().")
    for tenant, board in WORKDAY_COMPANIES:
        results[(tenant, board)] = {"ats": "workday", "status": "skipped"}

    # ── Patch file ─────────────────────────────────────────────────────────────
    needs_patch = any(v["status"] in ("fixed", "failed") for v in results.values()
                      if v["ats"] != "workday")
    if needs_patch:
        print("\nPatching job_search.py...")
        _patch_job_search_py(results)
        print("  Done.")

    # ── Summary ────────────────────────────────────────────────────────────────
    def _counts(ats: str):
        subset = [v for v in results.values() if v["ats"] == ats]
        ok = sum(1 for v in subset if v["status"] == "ok")
        fixed = sum(1 for v in subset if v["status"] == "fixed")
        failed = sum(1 for v in subset if v["status"] == "failed")
        skipped = sum(1 for v in subset if v["status"] == "skipped")
        return ok, fixed, failed, skipped

    gh_ok, gh_fix, gh_fail, _ = _counts("greenhouse")
    lv_ok, lv_fix, lv_fail, _ = _counts("lever")
    _, _, _, wd_skip = _counts("workday")

    print("\n" + "━" * 40)
    print("VALIDATION SUMMARY")
    print(f"  Greenhouse: {gh_ok} ok, {gh_fix} fixed, {gh_fail} failed")
    print(f"  Lever:      {lv_ok} ok, {lv_fix} fixed, {lv_fail} failed")
    print(f"  Workday:    {wd_skip} skipped (Cloudflare-protected — validated at runtime)")

    failed_entries = [(k, v) for k, v in results.items()
                      if v["status"] == "failed" and v["ats"] != "workday"]
    if failed_entries:
        print("\nUNCONFIRMED SLUGS — manual fix needed")
        print("━" * 40)
        for key, info in failed_entries:
            ats = info["ats"]
            if ats == "greenhouse":
                print(f"  {key}: check https://boards.greenhouse.io/{key}")
            else:
                print(f"  {key}: check https://jobs.lever.co/{key}")
    else:
        print("\nAll Greenhouse and Lever slugs confirmed ✅")


# ─── End slug validation helpers ──────────────────────────────────────────────


def fetch_json(url: str) -> dict:
    """Fetch URL and return parsed JSON. Returns {} on failure."""
    try:
        r = requests.get(url, headers=HEADERS, timeout=REQUEST_TIMEOUT)
        r.raise_for_status()
        return r.json()
    except Exception as e:
        print(f"  [warn] Could not fetch JSON from {url}: {e}")
        return {}


SWE_TITLE_KEYWORDS = [
    "engineer", "developer", "software", "sde", "swe", "backend", "frontend",
    "full stack", "fullstack", "infrastructure", "security engineer",
    "machine learning", "data engineer", "devops", "cloud engineer",
    "systems", "compiler", "firmware", "embedded", "mobile", "android", "ios",
]

NON_SWE_TITLE_KEYWORDS = [
    "account executive", "account manager", "sales", "recruiter", "recruiting",
    "finance", "legal", "marketing", "design", "customer success", "customer support",
    "operations", "analyst", "consultant", "business development", "partnership",
    "product manager", "program manager", "project manager", "hr ", "human resources",
]

def is_entry_level_title(title: str) -> bool:
    """Return True if the title is an entry-level SWE role (not senior/management, not non-SWE)."""
    t = title.lower()
    if any(kw in t for kw in SENIOR_TITLE_KEYWORDS):
        return False
    if any(kw in t for kw in NON_SWE_TITLE_KEYWORDS):
        return False
    return any(kw in t for kw in SWE_TITLE_KEYWORDS)


def fetch_greenhouse_jobs(companies: list[str]) -> list[dict]:
    """
    Fetch entry-level SWE jobs from Greenhouse company boards.
    The Greenhouse API returns the full job description in the `content` field,
    so no separate page scrape is needed for these jobs.
    """
    import html as html_module
    jobs = []
    for slug in companies:
        url = f"https://boards-api.greenhouse.io/v1/boards/{slug}/jobs?content=true"
        data = fetch_json(url)
        for job in data.get("jobs", []):
            title = job.get("title", "")
            if not is_entry_level_title(title):
                continue
            # extract plain text from HTML job description
            raw_html = html_module.unescape(job.get("content", ""))
            description = _parse_html_to_text(raw_html)[:8000]
            jobs.append({
                "company": job.get("company_name", slug),
                "role": title,
                "location": job.get("location", {}).get("name", ""),
                "url": job.get("absolute_url", ""),
                "posted": job.get("updated_at", ""),
                "description": description,
                "source": "greenhouse",
            })
        time.sleep(REQUEST_DELAY)
    return jobs


def fetch_lever_jobs(companies: list[str]) -> list[dict]:
    """
    Fetch entry-level SWE jobs from Lever ATS company boards.
    Lever returns full plain-text descriptions in the API response — no page scrape needed.
    """
    jobs = []
    for slug in companies:
        url = f"https://api.lever.co/v0/postings/{slug}?mode=json"
        postings = fetch_json(url)
        if isinstance(postings, dict):
            postings = postings.get("data", [])
        if not isinstance(postings, list):
            continue
        for posting in postings:
            title = posting.get("text", "")
            if not is_entry_level_title(title):
                continue
            cats = posting.get("categories", {})
            created_ms = posting.get("createdAt", 0)
            posted_iso = datetime.utcfromtimestamp(created_ms / 1000).isoformat() if created_ms else ""
            description = posting.get("descriptionPlain", "") or _parse_html_to_text(posting.get("description", ""))
            jobs.append({
                "company": posting.get("company", slug.title()),
                "role": title,
                "location": cats.get("location", ""),
                "url": posting.get("hostedUrl", ""),
                "posted": posted_iso,
                "description": description[:8000],
                "source": "lever",
            })
        time.sleep(REQUEST_DELAY)
    return jobs


def _find_workday_subdomain(tenant: str) -> int | None:
    """Find which wd1–wd5 subdomain a Workday tenant is on by checking DNS."""
    import socket
    for n in range(1, 6):
        host = f"{tenant}.wd{n}.myworkdayjobs.com"
        try:
            socket.getaddrinfo(host, 443)
            return n
        except OSError:
            pass
    return None


def fetch_workday_jobs(companies: list[tuple]) -> list[dict]:
    """
    Fetch entry-level SWE jobs from Workday ATS boards via Playwright (Cloudflare-protected).
    Workday's JSON search API requires a real browser session — we use Playwright to intercept
    the XHR response from the jobs endpoint.
    """
    if not PLAYWRIGHT_AVAILABLE:
        print("  [skip] Workday: Playwright not installed — run: pip install playwright && playwright install chromium")
        return []

    jobs = []
    from playwright.sync_api import sync_playwright

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(user_agent=(
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"
        ))

        for tenant, board in companies:
            n = _find_workday_subdomain(tenant)
            if not n:
                print(f"  [skip] Workday: DNS not found for {tenant}")
                continue
            base = f"https://{tenant}.wd{n}.myworkdayjobs.com"
            search_page = f"{base}/en-US/{board}?q=software+engineer"

            captured = []

            def _on_response(response):
                if "/jobs" in response.url and "wday/cxs" in response.url:
                    try:
                        data = response.json()
                        captured.append(data)
                    except Exception:
                        pass

            page = context.new_page()
            page.on("response", _on_response)
            try:
                page.goto(search_page, timeout=20000, wait_until="networkidle")
                page.wait_for_timeout(2000)
            except Exception as e:
                print(f"  [warn] Workday page load failed for {tenant}: {e}")
                page.close()
                time.sleep(REQUEST_DELAY)
                continue

            page.close()

            for data in captured:
                for posting in data.get("jobPostings", []):
                    title = posting.get("title", "")
                    if not is_entry_level_title(title):
                        continue
                    ext_path = posting.get("externalPath", "")
                    apply_url = f"{base}/{board}/job/{ext_path}" if ext_path else ""
                    jobs.append({
                        "company": tenant.title(),
                        "role": title,
                        "location": posting.get("locationsText", ""),
                        "url": apply_url,
                        "posted": "",
                        "description": "",
                        "source": "workday",
                    })
            time.sleep(REQUEST_DELAY)

        browser.close()
    return jobs


SYSTEM_PROMPT = """
You are a strict, expert job-fit evaluator. Your job is to score how well a specific
job posting matches a specific candidate profile. You return ONLY a JSON object.
No preamble. No explanation outside the JSON. No markdown. Just raw JSON.

Output format (always exactly this):
{"score": <0-100>, "reason": "<2-3 sentences explaining the score>", "verdict": "<STRONG_FIT|GOOD_FIT|NEAR_MISS|FLAGGED|EXCLUDED>", "flag": "<null or specific disqualifier reason>"}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
CANDIDATE PROFILE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Name: Shweta Patel
Location: Seattle, WA — open to relocate anywhere in the US
Work Authorization: F-1 visa, OPT starting September 14 2026 (does NOT need employer
sponsorship to start). STEM OPT extension eligible (24 months after initial OPT).
Exclude only if posting explicitly says: "no visa", "US citizens only",
"security clearance required", "must be authorized without sponsorship."
OPT itself does not require sponsorship — do not penalize for lack of H1B sponsorship.

Experience: 6+ years total
- Techblocks Consulting (Oct 2021–Sep 2024): Senior-level C#/.NET Core microservices,
  REST APIs, Azure cloud, SignalR, NUnit/xUnit, Azure DevOps CI/CD, Redis, CQRS,
  Repository Pattern, Agile. 500–1000 concurrent users. 40% API performance improvement.
- Meditab Software (Jul 2018–Oct 2021): ASP.NET MVC, .NET Core REST APIs, AWS IoT Core
  (MQTT pub/sub), Amazon Lex chatbot + Lambda, SQL Server T-SQL optimization,
  Firebase, SOLID principles, Windows Services, SOAP services.

Education: M.S. Computer Information Science, Harrisburg University (graduating Aug 2026,
STEM-designated, GPA 4.0). B.E. Computer Engineering (GPA 8.3/10).

Core Technical Skills (VERIFIED, HANDS-ON):
- Languages: C#, JavaScript, TypeScript, SQL. Python (personal projects — include when JD requires it).
- Frameworks: .NET Core, ASP.NET Core, ASP.NET MVC, Entity Framework, LINQ, SignalR
- Cloud Azure: App Service, Functions, Service Bus, Key Vault, Application Insights,
  Azure SQL, Azure DevOps (APPLICATION-LAYER only — not infra/networking)
- Cloud AWS: IoT Core, Lambda, Lex (real production usage)
- Cloud GCP: Pub/Sub, Dataflow, BigQuery (personal project)
- Databases: SQL Server, Azure SQL, Redis, BigQuery, Entity Framework
- DevOps: Azure DevOps CI/CD, GitHub Actions, Docker, Git, YAML pipelines
- Testing: NUnit, xUnit (90%+ coverage), Moq
- Patterns: CQRS, Repository Pattern, Microservices, REST APIs, Event-driven Pub/Sub,
  SOLID, Observer, Factory, Strategy
- Other: Swagger, Postman, IIS, Windows Services, SOAP

Personal Projects:
1. Queue Backlog Intelligence System (QBIS): C#/.NET 8, Azure Functions, Azure Service Bus,
   Azure Table Storage, React 18, Docker. Serverless monitoring system with 4 Azure Functions,
   10 REST endpoints, sub-45-second alert latency. Real production-quality architecture.
2. AI Job Search Automation Pipeline: Python 3.11, Google Gemini API, Groq API, Playwright,
   BeautifulSoup4, Greenhouse REST API, argparse CLI. Scrapes 3700+ job postings daily,
   LLM scoring, dual provider fallback. Live on GitHub, actively maintained.
   This counts as real applied AI/ML and Python experience.
3. COVID-19 Data Pipeline: Python, GCP Pub/Sub, Dataflow (Apache Beam), BigQuery.
   Streaming pipeline, 5-10M records, reduced processing from 30-60min to under 5min.

React.js: CONDITIONAL — only mention/count if the JD specifically requires React.
Python: CONDITIONAL — only count as primary if JD specifically requires it.

NOT skilled in (DO NOT credit these even if JD mentions them):
- Terraform / Infrastructure as Code
- Kubernetes / container orchestration
- IAM / Entra ID platform engineering
- Cloud networking (VPC/VNet, Private Link, DNS, firewalls)
- GCP infra (Cloud KMS, VPC — only used Pub/Sub, Dataflow, BigQuery)
- Embedded systems, firmware, FPGA, RTOS
- Angular (no proficiency — skip if core requirement)
- QA/test automation frameworks (Playwright for browser automation is NOT a test framework)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
HARD DISQUALIFIERS — SCORE 0 IMMEDIATELY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Return score=0, verdict=EXCLUDED if ANY of these are true:

1. CLEARANCE/CITIZENSHIP REQUIRED:
   Title or description contains: "Top Secret", "TS/SCI", "Secret Clearance",
   "Security Clearance", "Public Trust", "ITAR", "DoD", "US Citizen required",
   "US Citizenship Required", "Must be a US Person", "Cleared", "Poly"

2. NON-US LOCATION:
   Job is in UK, India, Canada, Europe, Australia, or any country outside the US.
   "Remote in UK", "London", "Toronto", "Bangalore" etc. -> Score 0.
   "Remote" alone with no country = flag as "verify location" but do not exclude.

3. WRONG ROLE TYPE (title-level only):
   Title contains: "firmware engineer", "embedded engineer", "fpga", "rtos",
   "avionics", "SDET", "QA engineer", "test automation engineer",
   "network engineer", "mainframe developer"

4. CONTRACT/PLACEMENT ARRANGEMENT:
   Description contains: "contract to hire", "w2 contract", "consulting engagement",
   "placed at client", "client site", "through our client", "on behalf of our client",
   "contingent", "6 month contract", "12 month contract"

5. EXPLICIT NO SPONSORSHIP:
   "No sponsorship", "will not sponsor", "no visa sponsorship available",
   "must be authorized to work without sponsorship", "citizens only"
   NOTE: silence on sponsorship is NOT a disqualifier. OPT does not require sponsorship.

6. SENIORITY MISMATCH:
   Title contains: "Senior", "Sr.", "Staff", "Principal", "Lead", "Director",
   "VP", "Head of", "Architect", "Distinguished", "Fellow"
   Target is entry-level, junior, mid-level, associate (0-5 years).

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
LEGITIMACY SCREEN — CAP SCORE AT 40 IF ANY MATCH
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Flag and cap score at maximum 40 if:
- Salary is clearly a data entry error (e.g., $680,000 for a junior role)
- Title says "Software Engineer" but skills list is "Customer Service, MS Office"
- 20+ completely unrelated technologies listed with no coherent focus
- Company name is vague with no web presence ("XYZ Solutions LLC")
- Non-tech organization posting generic SWE templates (community groups, NGOs)
- Known defense contractors (cap at 50, flag as "likely clearance required, verify before applying"):
  Lockheed Martin, Booz Allen Hamilton, Leidos, Raytheon, RTX, Peraton, L3Harris,
  SAIC, General Dynamics, Northrop Grumman, Sierra Nevada Corporation, Boeing Defense,
  BAE Systems
- Known aggregator/reseller posting (FetchJobs.co, similar): Score 0.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
OR-LIST RULE — CRITICAL, DO NOT SKIP
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

When a JD says "experience in Java, Python, C#, Go, or TypeScript":
- This IS fully satisfied by C# alone — do NOT penalize.
- Do NOT treat as partial credit or reduce the score.

EXCEPTION: If every responsibility example uses a specific language (e.g., "build Java
microservices", "maintain our Java codebase") and C# is only in the OR-list, then the
OR-list is COSMETIC — score the stack match low (30-50 range).

For LARGE, WELL-KNOWN companies (500+ employees, brand names): OR-lists are INTENTIONAL.
Score the OR-list as fully satisfied. Focus scoring on responsibilities.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SCORING RUBRIC
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Score 0: Hard disqualifier triggered
Score 10-30: Wrong stack, wrong domain, or legitimacy issue
Score 31-50: Some overlap but major gaps or red flags
Score 51-65: Decent match but Angular-heavy, unclear stack, or company/role concern
Score 66-79: Good fit — stack matches, responsibilities align, minor gaps
Score 80-89: Strong fit — C#/.NET primary or dominant OR-list, Azure/cloud experience matches
Score 90-100: Near-perfect — C#/.NET explicit primary, Azure matches what candidate built,
              seniority right, direct hire, domain is healthcare/fintech/enterprise

WHAT RAISES THE SCORE:
+ C#/.NET/.NET Core is explicitly the primary language
+ Azure: App Service, Functions, Service Bus, Key Vault, Application Insights, Azure SQL, Azure DevOps
+ AWS: IoT Core, Lambda
+ Microservices, REST APIs, distributed systems as core responsibilities
+ Healthcare, fintech, insurance, or enterprise software domain
+ Remote or hybrid work arrangement
+ Known H1B sponsor company
+ Role explicitly says entry/junior/associate/mid-level or 0-5 years
+ SignalR, Redis, Entity Framework, LINQ mentioned
+ CI/CD, Azure DevOps, GitHub Actions mentioned
+ CQRS, Repository Pattern, SOLID mentioned
+ LLM integration, AI tooling, data pipelines, or developer tooling mentioned

WHAT LOWERS THE SCORE (does not disqualify):
- Angular listed as core/required -> -15 points
- Terraform or Kubernetes as core requirement -> -10 points
- Primary stack is Java/Go/Python with C# as cosmetic OR-list -> -20 points
- On-site only (no remote/hybrid) -> -5 points
- No salary disclosed -> -3 points
- Small company (<50 employees) with no verifiable web presence -> -5 points
- Role requires 5+ years with no flexibility -> -5 points

ANGULAR RULE:
- Angular as CORE/REQUIRED -> -15 pts
- Angular as NICE-TO-HAVE/MENTIONED -> no penalty

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
KNOWN FALSE POSITIVE PATTERNS — SCORE 0 IMMEDIATELY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

- "Adtran Software Engineer" -> embedded C/C++ firmware, NOT C#/.NET
- "Peraton" any role -> defense contractor, almost always clearance required
- "DataAnnotation" -> gig-based AI labeling, not software engineering
- "FetchJobs.co" -> job board aggregator, not a real employer
- "Mon Sinistre Afrique", "Rotaract Club", "IndTech Calibration" -> non-tech orgs
- HackerRank SDET -> QA/test automation, not backend engineering
- Any Sophos role "Remote in UK" -> non-US, hard disqualifier
- Boeing "Software Engineer-Simulation" -> avionics/defense simulation

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
FINAL CHECK BEFORE SCORING
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Before returning your score, verify these 5 things:
1. Did any hard disqualifier trigger? -> Score 0 if yes
2. Is C#/.NET the actual core stack, or cosmetic in an OR-list?
3. Is the role type correct? (backend/full-stack SWE, NOT QA/infra/embedded/data eng)
4. Is the location actually in the US?
5. Are the responsibilities things she has actually done?
   (microservices, REST APIs, cloud services, backend systems)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
OUTPUT — RETURN ONLY THIS JSON, NOTHING ELSE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{"score": <0-100>, "reason": "<2-3 sentences>", "verdict": "<STRONG_FIT|GOOD_FIT|NEAR_MISS|FLAGGED|EXCLUDED>", "flag": "<null or disqualifier reason>"}
"""


def score_job(provider: str, client, profile: str, job: dict, description: str) -> dict:
    """
    Score job fit using the expert SYSTEM_PROMPT evaluator.
    Returns dict with score, reason, verdict, flag.
    """
    job_text = f"""CANDIDATE PROFILE:
{profile}

Job to evaluate:
Title: {job['role']}
Company: {job['company']}
Location: {job['location']}
Description: {description if description else '(no description available — score based on title only)'}"""

    try:
        text = call_llm(provider, client, job_text, max_tokens=512, system_prompt=SYSTEM_PROMPT)

        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)

        import json
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if match:
            data = json.loads(match.group())
            return {
                "score": int(data.get("score", 0)),
                "reason": str(data.get("reason", "")),
                "verdict": str(data.get("verdict", "")),
                "flag": str(data.get("flag") or ""),
            }
    except Exception as e:
        print(f"  [warn] Scoring failed for {job['company']} - {job['role']}: {e}")

    return {"score": 0, "reason": "scoring error", "verdict": "EXCLUDED", "flag": "parse_error"}


# ── LLM provider ──────────────────────────────────────────────────────────────

def call_llm(provider: str, client, prompt: str, max_tokens: int = 1024, system_prompt: str = None) -> str:
    """Unified LLM call — works with either Gemini or Groq client."""
    if provider == "gemini":
        from google.genai import types
        config = types.GenerateContentConfig(system_instruction=system_prompt) if system_prompt else None
        response = client.models.generate_content(model=GEMINI_MODEL, contents=prompt, config=config)
        return response.text.strip()
    else:
        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        messages.append({"role": "user", "content": prompt})
        response = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=messages,
            max_tokens=max_tokens,
        )
        return response.choices[0].message.content.strip()


def get_llm_client() -> tuple:
    """
    Auto-detect which LLM provider to use.
    Checks GEMINI_API_KEY first (primary), falls back to GROQ_API_KEY.
    Returns (provider_name, client_object).
    """
    gemini_key = os.environ.get("GEMINI_API_KEY", "")
    groq_key = os.environ.get("GROQ_API_KEY", "")

    if gemini_key:
        from google import genai
        print("Using Google Gemini API")
        return ("gemini", genai.Client(api_key=gemini_key))

    if groq_key:
        from groq import Groq
        print("Using Groq API (local fallback)")
        return ("groq", Groq(api_key=groq_key))

    print("Error: No API key found. Set one of:")
    print("  Primary:  export GEMINI_API_KEY='your-key'  # https://aistudio.google.com/app/apikey")
    print("  Fallback: export GROQ_API_KEY='your-key'    # https://console.groq.com/keys")
    sys.exit(1)


# ── Main ──────────────────────────────────────────────────────────────────────


def main():
    parser = argparse.ArgumentParser(description="Job Search Automation")
    parser.add_argument("--init", metavar="PROFILE_DOCX", help="Build profile_data.md from a .docx resume/profile")
    parser.add_argument("--update", metavar="FACT", help="Merge a new fact into profile_data.md")
    parser.add_argument("--show", action="store_true", help="Print the current profile_data.md")
    parser.add_argument("--today", action="store_true", help="Only include jobs posted within the last 24 hours")
    parser.add_argument("--validate-slugs", action="store_true",
                        help="Test all ATS slugs live, auto-fix variations, comment out failures")
    args = parser.parse_args()

    if args.validate_slugs:
        run_slug_validation()
        return

    if args.show:
        print(load_profile_doc())
        return

    if args.init:
        if not os.path.exists(args.init):
            print(f"Error: file not found: {args.init}")
            sys.exit(1)
        print(f"Reading profile from: {args.init}")
        raw_profile = read_docx(args.init)
        if not raw_profile.strip():
            print("Error: could not extract text from the .docx file.")
            sys.exit(1)
        print(f"  Profile loaded ({len(raw_profile)} chars)")

        provider, client = get_llm_client()
        print("\nSummarizing profile...")
        profile_text = summarize_profile(provider, client, raw_profile)
        save_profile_doc(profile_text)

        print("\n" + "─" * 60)
        print(f"Saved profile doc to: {PROFILE_DOC_PATH}")
        print("─" * 60)
        print(profile_text)
        print("─" * 60)
        return

    if args.update:
        provider, client = get_llm_client()
        current_doc = load_profile_doc()
        print("Updating profile doc with new fact...")
        updated_doc = update_profile_doc(provider, client, current_doc, args.update)
        save_profile_doc(updated_doc)

        print("\n" + "─" * 60)
        print(f"Updated profile doc saved to: {PROFILE_DOC_PATH}")
        print("─" * 60)
        print(updated_doc)
        print("─" * 60)
        return

    # ── Normal run: use the existing profile doc directly ──
    if not PLAYWRIGHT_AVAILABLE:
        print("Tip: install Playwright to scrape JS-rendered job pages:")
        print("     pip install playwright && playwright install chromium")
        print()

    profile_text = load_profile_doc()
    provider, client = get_llm_client()

    # ── Fetch jobs from all sources ──
    all_jobs = []
    for source in GITHUB_SOURCES:
        print(f"\nFetching jobs from {source['name']}...")
        md = fetch_text(source["raw_url"])
        if not md:
            print(f"  [skip] Could not fetch {source['name']}")
            continue
        jobs = parse_jobs(md, source.get("format", "markdown-html"))
        print(f"  Found {len(jobs)} open positions")
        for job in jobs:
            job["source"] = source["name"]
        all_jobs.extend(jobs)
        time.sleep(REQUEST_DELAY)

    print(f"\nFetching jobs from Greenhouse ({len(GREENHOUSE_COMPANIES)} companies)...")
    gh_jobs = fetch_greenhouse_jobs(GREENHOUSE_COMPANIES)
    print(f"  Found {len(gh_jobs)} entry-level positions across Greenhouse boards")
    all_jobs.extend(gh_jobs)

    print(f"\nFetching jobs from Lever ({len(LEVER_COMPANIES)} companies)...")
    lv_jobs = fetch_lever_jobs(LEVER_COMPANIES)
    print(f"  Found {len(lv_jobs)} entry-level positions across Lever boards")
    all_jobs.extend(lv_jobs)

    print(f"\nFetching jobs from Workday ({len(WORKDAY_COMPANIES)} companies)...")
    wd_jobs = fetch_workday_jobs(WORKDAY_COMPANIES)
    print(f"  Found {len(wd_jobs)} entry-level positions across Workday boards")
    all_jobs.extend(wd_jobs)

    # deduplicate by URL
    seen_urls = set()
    unique_jobs = []
    for job in all_jobs:
        if job["url"] not in seen_urls:
            seen_urls.add(job["url"])
            unique_jobs.append(job)

    print(f"\nTotal unique jobs found: {len(unique_jobs)}")

    if args.today:
        before = len(unique_jobs)
        unique_jobs = [j for j in unique_jobs if is_posted_today(j.get("posted", ""))]
        print(f"  Filtered to jobs posted in the last 24 hours: {len(unique_jobs)} (of {before})")

    print(f"\nTotal unique jobs to process: {len(unique_jobs)}")
    if len(unique_jobs) > MAX_JOBS:
        print(f"Capping at {MAX_JOBS} jobs. Edit MAX_JOBS in the script to change this.")
        unique_jobs = unique_jobs[:MAX_JOBS]

    # ── Score each job ──
    results = []

    for i, job in enumerate(unique_jobs, 1):
        print(f"\n[{i}/{len(unique_jobs)}] {job['company']} — {job['role']}")
        # Greenhouse jobs already carry the description from the API — skip page scrape
        if job.get("description"):
            print(f"  Description pre-loaded from Greenhouse API")
            description = job["description"]
        else:
            print(f"  Fetching job description from: {job['url']}")
            description = extract_job_description(job["url"])
            time.sleep(REQUEST_DELAY)

        print(f"  Scoring...")
        scored = score_job(provider, client, profile_text, job, description)
        job["score"] = scored["score"]
        job["reason"] = scored["reason"]
        job["verdict"] = scored.get("verdict", "")
        job["flag"] = scored.get("flag") or ""
        print(f"  Score: {job['score']}/100 [{job['verdict']}] — {job['reason']}")
        if job["flag"]:
            print(f"  Flag: {job['flag']}")

        if job["score"] >= MIN_SCORE and job["verdict"] != "EXCLUDED":
            results.append(job)

    # ── Write CSV ──
    results.sort(key=lambda x: x["score"], reverse=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = f"jobs_{timestamp}.csv"

    fieldnames = ["score", "verdict", "flag", "company", "role", "location", "posted", "reason", "source", "url"]
    with open(output_file, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for row in results:
            writer.writerow({k: row.get(k, "") for k in fieldnames})

    print(f"\n{'='*60}")
    print(f"Done! {len(results)} matching jobs (score >= {MIN_SCORE}) saved to:")
    print(f"  {os.path.abspath(output_file)}")
    print(f"{'='*60}")
    if not results:
        print("No jobs met the minimum score threshold. Try lowering MIN_SCORE in the script.")


if __name__ == "__main__":
    main()
